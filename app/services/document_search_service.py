"""
文档搜索服务
提供Word文档的搜索和处理功能
"""

import os
import json
import time
import hashlib
from typing import List, Dict
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
import jieba
import re
import subprocess

from app.models import DocumentBlock
from app.services.keyword_extractor import KeywordExtractor
from config import SEARCH_CONFIG


class DocumentSearchService:
    """
    文档搜索服务
    提供基于关键词的Word文档搜索功能
    """
    
    def __init__(self, folder_path: str, enhanced_context: bool = True):
        """
        初始化文档搜索服务
        
        Args:
            folder_path: Word文档所在的文件夹路径
            enhanced_context: 是否使用增强的上下文抓取模式
        """
        self.folder_path = folder_path
        self._cache = {}  # 搜索结果缓存
        self._doc_cache = {}  # 文档对象缓存
        self.results = []  # 最新搜索结果
        self.keyword_extractor = KeywordExtractor()
        self.enhanced_context = enhanced_context  # 上下文模式
        
    @lru_cache(maxsize=SEARCH_CONFIG["cache_size"])
    def _get_document(self, file_path: str) -> Document:
        """
        获取文档对象（带缓存）
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            Document: Word文档对象
        """
        return Document(file_path)
    
    def _get_cache_key(self, keywords: List[str]) -> str:
        """
        生成搜索缓存键
        
        Args:
            keywords: 关键词列表
            
        Returns:
            str: 缓存键
        """
        return hashlib.md5(','.join(sorted(keywords)).encode('utf-8')).hexdigest()
    
    @staticmethod
    def is_block_start(text: str) -> bool:
        # 临时调试：每一段都当作文段
        return True
    
    def get_heading_context(self, doc: Document, para_index: int) -> str:
        """
        获取段落的上下文标题信息（改进版）
        支持多种标题格式识别和层级分析
        
        Args:
            doc: Word文档对象
            para_index: 段落索引
            
        Returns:
            str: 上下文标题字符串
        """
        context_parts = []
        paragraphs = doc.paragraphs
        
        # 收集上下文候选项
        context_candidates = []
        
        # 向上查找上下文信息（最多查找20个段落）
        search_range = min(para_index + 1, SEARCH_CONFIG["max_context_search"])
        
        for i in range(para_index, max(-1, para_index - search_range), -1):
            if i < len(paragraphs):
                para = paragraphs[i]
                text = self.clean_text(para.text)
                
                if not text:  # 跳过空段落
                    continue
                
                # 多种标题识别方法
                title_score = self._calculate_title_score(para, text)
                
                if title_score > 0:
                    context_candidates.append({
                        'index': i,
                        'text': text,
                        'score': title_score,
                        'level': self._determine_title_level(text)
                    })
        
        # 按层级和得分排序，构建层级上下文
        if context_candidates:
            # 按层级从低到高排序（层级越低越重要）
            context_candidates.sort(key=lambda x: (x['level'], -x['score']))
            
            # 构建层级关系
            used_levels = set()
            for candidate in context_candidates:
                level = candidate['level']
                # 避免同一层级重复，选择得分最高的
                if level not in used_levels:
                    context_parts.append(candidate['text'])
                    used_levels.add(level)
                    # 最多保留3个层级
                    if len(context_parts) >= 3:
                        break
        
        # 如果没有找到标题，尝试查找重要段落
        if not context_parts:
            important_para = self._find_important_paragraph(paragraphs, para_index)
            if important_para:
                context_parts.append(important_para)
        
        return ' > '.join(reversed(context_parts)) if context_parts else "文档内容"
    
    def _calculate_title_score(self, para, text: str) -> float:
        """
        计算段落作为标题的可能性得分
        
        Args:
            para: Word段落对象
            text: 段落文本
            
        Returns:
            float: 标题得分 (0-1)
        """
        score = 0.0
        
        # 1. Word样式检查（权重：0.4）
        if para.style.name.startswith('Heading'):
            score += 0.4
        elif '标题' in para.style.name or 'Title' in para.style.name:
            score += 0.3
        
        # 2. 格式特征检查（权重：0.3）
        # 数字编号格式
        if re.match(r'^[0-9]+(\.[0-9]+)*[、.]?\s*[^\s]', text):
            score += 0.15
        # 中文编号格式
        elif re.match(r'^[一二三四五六七八九十]+[、.]?\s*[^\s]', text):
            score += 0.15
        # 英文编号格式
        elif re.match(r'^[A-Z]+[、.]?\s*[^\s]', text):
            score += 0.1
        # 罗马数字格式
        elif re.match(r'^[IVX]+[、.]?\s*[^\s]', text):
            score += 0.1
        
        # 3. 内容特征检查（权重：0.2）
        # 长度适中（10-100字符）
        if 10 <= len(text) <= 100:
            score += 0.05
        # 包含关键词
        title_keywords = ['概述', '简介', '背景', '目标', '方法', '步骤', '流程', '配置', 
                         '安装', '部署', '管理', '监控', '故障', '排查', '总结', '结论',
                         'overview', 'introduction', 'background', 'objective', 'method',
                         'process', 'configuration', 'installation', 'deployment', 
                         'management', 'monitoring', 'troubleshooting', 'summary']
        if any(keyword in text.lower() for keyword in title_keywords):
            score += 0.1
        
        # 4. 字体格式检查（权重：0.1）
        # 注：这里简化处理，实际可以检查字体大小、加粗等
        if len(text) < 50 and not text.endswith('。'):  # 短句且不以句号结尾
            score += 0.05
        
        return min(score, 1.0)
    
    def _determine_title_level(self, text: str) -> int:
        """
        确定标题层级
        
        Args:
            text: 标题文本
            
        Returns:
            int: 标题层级 (1-6, 数字越小层级越高)
        """
        # 数字编号层级判断
        number_match = re.match(r'^([0-9]+(?:\.[0-9]+)*)', text)
        if number_match:
            levels = number_match.group(1).count('.') + 1
            return min(levels, 6)
        
        # 中文编号层级判断
        if re.match(r'^[一二三四五六七八九十]+', text):
            return 1
        
        # 英文编号层级判断
        if re.match(r'^[A-Z]+[、.]', text):
            return 2
        
        # 默认为较低层级
        return 3
    
    def _find_important_paragraph(self, paragraphs: List, para_index: int) -> str:
        """
        查找重要段落作为上下文
        
        Args:
            paragraphs: 段落列表
            para_index: 当前段落索引
            
        Returns:
            str: 重要段落文本，如果没找到返回None
        """
        # 向上查找最近的非空、有意义的段落
        for i in range(para_index - 1, max(-1, para_index - 10), -1):
            if i < len(paragraphs) and i >= 0:
                text = self.clean_text(paragraphs[i].text)
                
                # 过滤掉过短或过长的段落
                if 20 <= len(text) <= 200:
                    # 确保不是列表项或代码
                    if not re.match(r'^[-*•]\s', text) and not re.match(r'^\s*\d+[.)]\s', text):
                        # 截取前50个字符作为上下文
                        return text[:50] + ('...' if len(text) > 50 else '')
        
        return None
    
    def _get_paragraph_context_enhanced(self, doc: Document, para_index: int, context_range: int = None) -> str:
        """
        获取增强的段落上下文（包括前后段落）
        
        Args:
            doc: Word文档对象
            para_index: 段落索引
            context_range: 上下文范围，如果为None则使用配置文件的设置
            
        Returns:
            str: 增强的上下文信息
        """
        if context_range is None:
            context_range = SEARCH_CONFIG["context_range"]
            
        paragraphs = doc.paragraphs
        context_parts = []
        
        # 获取主上下文（标题）
        main_context = self.get_heading_context(doc, para_index)
        if main_context and main_context != "文档内容":
            context_parts.append(main_context)
        
        # 获取相邻段落上下文
        adjacent_contexts = []
        
        # 前面的段落
        for i in range(max(0, para_index - context_range), para_index):
            if i < len(paragraphs):
                text = self.clean_text(paragraphs[i].text)
                if text and len(text) > 10:
                    # 取段落的前30个字符
                    preview = text[:30] + ('...' if len(text) > 30 else '')
                    adjacent_contexts.append(f"前文: {preview}")
        
        # 后面的段落
        for i in range(para_index + 1, min(len(paragraphs), para_index + context_range + 1)):
            text = self.clean_text(paragraphs[i].text)
            if text and len(text) > 10:
                # 取段落的前30个字符
                preview = text[:30] + ('...' if len(text) > 30 else '')
                adjacent_contexts.append(f"后文: {preview}")
        
        # 组合上下文
        if adjacent_contexts:
            context_parts.append(' | '.join(adjacent_contexts))
        
        return ' >> '.join(context_parts) if context_parts else "文档内容"
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        清理文本内容，去除多余字符和格式
        
        Args:
            text: 原始文本
            
        Returns:
            str: 清理后的文本
        """
        if not text:
            return ""
            
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除特殊字符，保留中文、英文、数字和常见标点
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,，。：:；;？！?！]', '', text)
        
        return text.strip()
    
    def _process_file(self, filename: str, keywords: List[str]) -> List[DocumentBlock]:
        """
        处理单个文档文件，提取匹配的文档块
        
        Args:
            filename: 文件名
            keywords: 搜索关键词列表
            
        Returns:
            List[DocumentBlock]: 匹配的文档块列表
        """
        # 跳过非doc/docx文件和临时文件
        if not filename.endswith('.docx') or filename.startswith('~$'):
            return []

        file_path = os.path.join(self.folder_path, filename)
        blocks = []
        try:
            if filename.endswith('.docx'):
                doc = self._get_document(file_path)
                paragraphs = doc.paragraphs
                n = len(paragraphs)
                i = 0
                while i < n:
                    text = self.clean_text(paragraphs[i].text)
                    if self.is_block_start(text):
                        block_start = i
                        block_texts = [text]
                        block_indices = [i]
                        i += 1
                        while i < n and not self.is_block_start(self.clean_text(paragraphs[i].text)):
                            current_text = self.clean_text(paragraphs[i].text)
                            if current_text:
                                block_texts.append(current_text)
                                block_indices.append(i)
                            i += 1
                        block_full_text = '\n'.join(block_texts)
                        found_keywords = [kw for kw in keywords if kw in block_full_text]
                        print(f"【调试】文段内容: {block_full_text[:50]}... 匹配到关键词: {found_keywords}")
                        if found_keywords:
                            if self.enhanced_context:
                                context = self._get_paragraph_context_enhanced(doc, block_start)
                            else:
                                context = self.get_heading_context(doc, block_start)
                            block = DocumentBlock(
                                filename=filename,
                                content=block_full_text,
                                context=context,
                                start_index=block_start,
                                keywords=found_keywords
                            )
                            blocks.append(block)
                    else:
                        i += 1
            return blocks
        except Exception as e:
            print(f"处理文件 {filename} 时出错: {str(e)}")
            return []
    
    def _calculate_similarity(self, block: DocumentBlock, keywords: List[str]) -> float:
        """
        计算文档块与关键词的相似度
        
        Args:
            block: 文档块对象
            keywords: 关键词列表
            
        Returns:
            float: 相似度分数 (0-1)
        """
        # 使用jieba分词
        block_words = set(jieba.lcut(block.content))
        keyword_set = set(keywords)
        
        # 计算Jaccard相似度
        intersection = len(keyword_set & block_words)
        union = len(keyword_set | block_words)
        
        return intersection / union if union > 0 else 0.0
    
    def search_documents(self, keywords: List[str]) -> List[Dict]:
        """
        在文档文件夹中搜索匹配关键词的内容
        
        Args:
            keywords: 搜索关键词列表
            
        Returns:
            List[Dict]: 搜索结果列表，按相似度排序
        """
        if not keywords:
            return []
            
        # 检查缓存
        cache_key = self._get_cache_key(keywords)
        if cache_key in self._cache:
            self.results = self._cache[cache_key]
            return self.results
        
        start_time = time.time()
        print(f"\n开始搜索文档，关键词: {', '.join(keywords)}")
        
        # 获取文件夹中的所有文件
        try:
            filenames = os.listdir(self.folder_path)
        except FileNotFoundError:
            print(f"错误：文档文件夹 '{self.folder_path}' 不存在")
            return []
        
        all_blocks = []
        
        # 使用线程池并行处理文件
        with ThreadPoolExecutor(max_workers=SEARCH_CONFIG["max_workers"]) as executor:
            # 提交任务
            future_to_file = {
                executor.submit(self._process_file, filename, keywords): filename
                for filename in filenames
            }
            
            # 收集结果
            for future in as_completed(future_to_file):
                filename = future_to_file[future]
                try:
                    blocks = future.result()
                    all_blocks.extend(blocks)
                except Exception as e:
                    print(f"处理文件 {filename} 时出错: {str(e)}")
        
        # 计算相似度并排序
        for block in all_blocks:
            block.similarity_score = self._calculate_similarity(block, keywords)
        
        # 按相似度排序并取前N个结果
        top_blocks = sorted(
            all_blocks, 
            key=lambda x: x.similarity_score, 
            reverse=True
        )[:SEARCH_CONFIG["max_results"]]
        
        # 转换为字典格式
        results = [block.to_dict() for block in top_blocks]
        
        # 计算耗时
        end_time = time.time()
        print(f"文档搜索完成，耗时: {end_time - start_time:.2f}秒，找到 {len(results)} 条结果")
        
        # 缓存结果
        self._cache[cache_key] = results
        self.results = results
        
        return results
    
    def save_results(self, output_file: str) -> None:
        """
        将搜索结果保存为JSON文件
        
        Args:
            output_file: 输出文件路径
        """
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(self.results, f, ensure_ascii=False, indent=2)
            print(f"搜索结果已保存到: {output_file}")
        except Exception as e:
            print(f"保存搜索结果时出错: {str(e)}")
    
    def clear_cache(self) -> None:
        """清除搜索缓存"""
        self._cache.clear()
        print("搜索缓存已清除")
    
    def get_cache_info(self) -> Dict:
        """
        获取缓存信息
        
        Returns:
            Dict: 缓存统计信息
        """
        return {
            "cache_size": len(self._cache),
            "max_cache_size": SEARCH_CONFIG["cache_size"],
            "cached_searches": list(self._cache.keys())
        } 